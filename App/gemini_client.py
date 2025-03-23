import os
import re
import json
from dotenv import load_dotenv
import spacy
from spacy import displacy
import re

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.exceptions import OutputParserException

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

load_dotenv()

nlp = spacy.load("en_core_web_sm")

def extract_data_from_resume(resume_text):
    doc = nlp(resume_text)
    
    name = None
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            name = ent.text
            break
        
    email = re.search(r'\S+@\S+', resume_text)
    email = email.group(0) if email else "N/A"
    
    phone = re.search(r'\+?\(?\d{1,3}\)?[-.\s]?\(?\d{1,3}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,4}', resume_text)
    phone = phone.group(0) if phone else "N/A"

    return name, email, phone

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id", r_id
    )
    new_run = OxmlElement("w:r")
    r_text = OxmlElement("w:t")
    r_text.text = text
    new_run.append(r_text)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)

class GeminiClient:
    def __init__(self):
        self.llm = ChatGoogleGenerativeAI(
            google_api_key=os.getenv("GOOGLE_API_KEY"),
            temperature=0,
            model="models/gemini-2.0-pro-exp-02-05"
        )

    def extract_jobs(self, cleaned_text):
        prompt_extract = PromptTemplate.from_template(
            """
            ### SCRAPED TEXT FROM WEBSITE:
            {page_data}

            ### INSTRUCTION:
            The scraped text is from a careers page.
            Extract job postings and return them in valid JSON with:
            role, experience, skills, description.
            """
        )
        chain_extract = prompt_extract | self.llm
        res = chain_extract.invoke({"page_data": cleaned_text})
        raw = res.content.strip()

        # Remove numeric keys + trailing commas, fix quotes
        raw = re.sub(r"\n?\s*\d+\s*:", "", raw)
        raw = re.sub(r",(\s*[\]}])", r"\1", raw)
        raw = raw.replace("“", "\"").replace("”", "\"")
        raw = raw.replace("‘", "'").replace("’", "'")

        try:
            json_parser = JsonOutputParser()
            parsed = json_parser.parse(raw)
            return parsed if isinstance(parsed, list) else [parsed]
        except OutputParserException as e:
            raise OutputParserException(f"❌ Failed to parse JSON:\n\n{raw}\n\nError: {e}")

    def write_mail(self, job_description, links):
        prompt_email = PromptTemplate.from_template(
            f"""
            ### TASK:
            Write a professional cold email for the job holder.

            Use the following:
            - Job Description: {job_description}
            - Portfolio links: {links}
            
            start with Dear Hiring Team if there is no explicit name mentioned in the job description.
            The mail itself sill consist of two short paragraphs.
            the framework utilized.
            
            acknowledge your hiring teams request for your resume
            reiterate your interest in the specific role you would like to be considered for
            ask if they need additional information to keep the process moving forward.
            
            keep it to less then 100 words 
            
            """
        )
        chain_email = prompt_email | self.llm
        res = chain_email.invoke({"job_desc": str(job_description), "link_list": links})
        return res.content

    def write_cover_letter(self, resume, job_description, links):
        prompt_cover = PromptTemplate.from_template(
        f"""
        ### TASK:
        Write a professional cover letter for the resume holder.

        Use the following:
        - Resume: {resume}
        - Job Description: {job_description}
        - Portfolio links: {links}

        Start with Dear Hiring Team if there is no explicit name mentioned in the job description.
        Focus on qualifications, fit, and enthusiasm. Make it tailored, clear, and confident.
        We will use the reason, anecdote, connection model (RAC for short).
        
        The cover letter answers the question: why should we consider you.

        Cover letters should not be unfacilitated. They should also be error free.

        The length should be 200 words max. 
        There should be an introductory paragraph, an RAC paragraph where you detail why you would be a good fit, and a short closing paragraph asking for an interview and contact info.

        For the introductory paragraph:

        State the position you are interested in, any key advocates you have at the organization (this should be an input parameter), an expression of belief that you can add value to the employer in “the following ways taken from the resume” which will lead you into the RAC paragraphs. I include a reason or motivation that explains why this position and why now. 

        Don't write too much here since It might go to a potential future boss, keep it lean and mean. 

        For the next paragraph:

        State a three noteworthy skills or attributes (Reason from RAC) that would make you appealing to the employer. Mostly soft skills (ex leadership:). The next sentence should provide an anecdote that shows why the employer finds your reason true or important.
        then go on to the next skill and repeat until you have gone through all the three skills. 

        The anecdote (sentence that comes after the reason)can take a few forms: 

        A brief summary of a bullet point on your resume that illustrates the skill or attribute (in a more natural language than what appears on the resume. It is proof that you possess a skill or attribute and your reason is true).
        A story about earning an accolade or award and demonstrating proficiency at the reason.( It is proof that you possess a skill or attribute and your reason is true)
        Information learned perhaps from online research of someone you spoke to in an informal meeting reinforces the importance of that reason within their organization and similar roles. (proves that your reason is important and matters to the employer.)
        
        keep the paragraph short and to the point as possible.

        The concluding paragraph:
        This should be a short closing paragraph that asks for an interview and contact info and connects your reasons to the employers values.
        If the reasons benefit the employer is not immediately obvious. Finish with one sentence that connects your reason to something the employer cares about (for example, will your skills increase sales?). This is the sentence where you demonstrate that you understand the job you’re applying for and appreciate the employer. The employer is the ultimate audience of this cover letter.  
 
        end with sincerely, name of applicant
        
        ### COVER LETTER (NO PREAMBLE):
        """
        )
        chain_cover = prompt_cover | self.llm
        res = chain_cover.invoke({
            "resume_data": resume,
            "job_desc": job_description,
            "link_list": links
        })
        return res.content

    def save_cover_letter(self, content, filename="Cover_Letter.docx"):
        name, email, phone = extract_data_from_resume(content)
        doc = Document()
        
        if name:
            # Add the heading with the name
            heading = doc.add_paragraph()
            run = heading.add_run(name)  # Add the name text as a run
            run.bold = True  # Make the text bold
            run.font.size = Pt(18)  # Set the font size
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the heading
        else:
            contact_info = f"Contact Information: {phone if phone else 'N/A'} | {email if email else 'N/A'}"
            heading = doc.add_paragraph()
            run = heading.add_run(contact_info)
            run.bold = True  # Make the contact info bold
            run.font.size = Pt(14)  # Set a different font size for the contact info
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the heading

        #heading.add_run(name).bold = True
        #run = heading.add_run(name)  # Add the name text as a run
        #run.bold = True  # Make it bold
        #run.font.size = Pt(18)
        #heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact = doc.add_paragraph()
        add_hyperlink(contact, phone, f"tel:{phone}")
        contact.add_run(" || ")
        add_hyperlink(contact, email, f"mailto:{email}")
        #contact.add_run(" || Auburn Hills, MI")
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

        #doc.add_paragraph()
        line = doc.add_paragraph()
        line.add_run("──────────────────────────────────────────────────────").bold = True
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER

        #doc.add_paragraph("\nSeptember 28, 2024")

        recipient = doc.add_paragraph()
        #recipient.add_run("Mr. Jeff Bennett\n").bold = True
        #recipient.add_run("Talent Acquisition Senior Manager, Deloitte LLP\n")
        #recipient.add_run("1001 Woodward,\n")
        #recipient.add_run("MI, 48226-1904\n")

        #doc.add_paragraph("\nDear Sir,\n")

        paragraphs = content.split("\n\n")
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph).style.font.size = Pt(12)

        #doc.add_paragraph("\nSincerely,")
        #doc.add_paragraph("Pradhum Niroula")

        doc.save(filename)
        return filename
